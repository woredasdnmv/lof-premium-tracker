# -*- coding: utf-8 -*-
"""将项目 Markdown 文档转换为 .docx 和 .pdf 格式"""
import os
import re
import sys

# 工作目录
PROJECT_DIR = r"D:\开发\金快查\在线版\lof-premium-tracker-main\lof-premium-tracker-main"
OUTPUT_DIR  = os.path.join(PROJECT_DIR, "docs", "exports")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 要导出的文件
FILES = [
    ("README.md", "金快查-项目说明"),
    ("docs/TECH.md", "金快查-技术文档"),
    ("docs/DEVELOPMENT.md", "金快查-开发指南"),
]


def read_md(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def md_to_text_blocks(md: str) -> list:
    """
    将 Markdown 解析为 (type, text) 块列表。
    type: heading1, heading2, heading3, paragraph, code, hr, list_item, table
    """
    lines = md.split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(code_lines)))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}\s*$", line.strip()):
            blocks.append(("hr", ""))
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,3})\s+(.+)", line)
        if m:
            level = len(m.group(1))
            tag = {1: "heading1", 2: "heading2", 3: "heading3"}[level]
            blocks.append((tag, m.group(2).strip()))
            i += 1
            continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", table_lines))
            continue

        # List item
        if re.match(r"^\s*[-*]\s+", line) or re.match(r"^\s*\d+\.\s+", line):
            blocks.append(("list_item", strip_md_inline(line.strip())))
            i += 1
            continue

        # Blockquote
        if line.strip().startswith(">"):
            blocks.append(("paragraph", strip_md_inline(line.strip()[1:].strip())))
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Paragraph
        blocks.append(("paragraph", strip_md_inline(line.strip())))
        i += 1

    return blocks


def strip_md_inline(text: str) -> str:
    """移除行内 Markdown 标记，保留纯文本"""
    # 图片
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    # 链接 → 保留文本
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # 粗体/斜体
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    # 行内代码
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # HTML 标签
    text = re.sub(r"<[^>]+>", "", text)
    # 移除多余的空白
    text = text.strip()
    return text


def parse_table(text: str) -> list:
    """解析 Markdown 表格，返回行列表"""
    rows = []
    for line in text:
        if re.match(r"^[\s|:-]+$", line):  # 分隔行
            continue
        cells = [cell.strip() for cell in line.strip().split("|")]
        cells = [c for c in cells if c]  # 去掉首尾空元素
        if cells:
            rows.append([strip_md_inline(c) for c in cells])
    return rows


def export_docx(md: str, output_path: str, title: str) -> None:
    """将 Markdown 导出为 Word 文档"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    blocks = md_to_text_blocks(md)

    for btype, btext in blocks:
        if btype == "heading1":
            h = doc.add_heading(btext, level=1)
            for run in h.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

        elif btype == "heading2":
            h = doc.add_heading(btext, level=2)
            for run in h.runs:
                run.font.size = Pt(15)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        elif btype == "heading3":
            h = doc.add_heading(btext, level=3)
            for run in h.runs:
                run.font.size = Pt(13)

        elif btype == "code":
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run(btext)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)

        elif btype == "hr":
            doc.add_paragraph("─" * 60)

        elif btype == "table":
            rows = parse_table(btext)
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]), style="Light Grid Accent 1")
            for ri, row_data in enumerate(rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < len(table.rows[ri].cells):
                        cell = table.rows[ri].cells[ci]
                        cell.text = cell_text
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)
            doc.add_paragraph()  # 表格后空行

        elif btype == "list_item":
            p = doc.add_paragraph(btext, style="List Bullet")

        elif btype == "paragraph":
            if btext:
                doc.add_paragraph(btext)

    doc.save(output_path)
    print(f"  [OK] .docx → {output_path}")


def export_pdf(md: str, output_path: str, title: str) -> None:
    """将 Markdown 导出为 PDF（使用 fpdf2）"""
    from fpdf import FPDF

    # 字体路径
    FONT_PATH  = "C:/Windows/Fonts/msyh.ttc"
    FONT_BOLD  = "C:/Windows/Fonts/msyhbd.ttc"

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.add_font("msyh", "", FONT_PATH, uni=True)
    pdf.add_font("msyh", "B", FONT_BOLD, uni=True)

    blocks = md_to_text_blocks(md)

    for btype, btext in blocks:
        if btype == "heading1":
            pdf.set_font("msyh", "B", 18)
            pdf.set_text_color(26, 86, 219)
            pdf.ln(6)
            pdf.multi_cell(0, 10, btext)
            pdf.set_draw_color(26, 86, 219)
            pdf.set_line_width(0.5)
            y = pdf.get_y()
            pdf.line(pdf.l_margin, y + 1, pdf.w - pdf.r_margin, y + 1)
            pdf.ln(6)

        elif btype == "heading2":
            pdf.set_font("msyh", "B", 14)
            pdf.set_text_color(44, 62, 80)
            pdf.ln(4)
            pdf.multi_cell(0, 8, btext)
            pdf.ln(3)

        elif btype == "heading3":
            pdf.set_font("msyh", "B", 12)
            pdf.set_text_color(52, 73, 94)
            pdf.ln(3)
            pdf.multi_cell(0, 7, btext)
            pdf.ln(2)

        elif btype == "code":
            pdf.set_font("msyh", "", 8)
            pdf.set_text_color(51, 51, 51)
            pdf.set_fill_color(245, 245, 245)
            for cline in btext.split("\n"):
                pdf.set_x(pdf.l_margin + 5)
                pdf.cell(0, 5, cline, new_x="LMARGIN", new_y="NEXT", fill=True)
            pdf.ln(3)

        elif btype == "hr":
            pdf.set_draw_color(200, 200, 200)
            y = pdf.get_y()
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(5)

        elif btype == "table":
            rows = parse_table(btext)
            if not rows:
                continue
            col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / len(rows[0])
            pdf.set_font("msyh", "", 8)
            # Header
            pdf.set_fill_color(26, 86, 219)
            pdf.set_text_color(255, 255, 255)
            for cell_text in rows[0]:
                pdf.cell(col_w, 8, cell_text, border=1, fill=True, new_x="RIGHT", new_y="LAST")
            pdf.ln()
            # Rows
            pdf.set_text_color(51, 51, 51)
            for row in rows[1:]:
                pdf.set_fill_color(249, 249, 249) if rows.index(row) % 2 == 1 else pdf.set_fill_color(255, 255, 255)
                for cell_text in row:
                    pdf.cell(col_w, 7, cell_text, border=1, fill=True, new_x="RIGHT", new_y="LAST")
                pdf.ln()
            pdf.set_text_color(51, 51, 51)
            pdf.ln(4)

        elif btype == "list_item":
            pdf.set_font("msyh", "", 10)
            pdf.set_x(pdf.l_margin + 5)
            pdf.cell(5, 6, "-", new_x="RIGHT")
            pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin - 10, 6, btext)
            pdf.ln(1)

        elif btype == "paragraph":
            if btext:
                pdf.set_font("msyh", "", 10)
                pdf.set_text_color(51, 51, 51)
                pdf.multi_cell(0, 6, btext)
                pdf.ln(2)

    pdf.output(output_path)
    print(f"  [OK] .pdf  → {output_path}")


def main():
    print("=" * 60)
    print("  金快查 — 项目文档导出工具")
    print("=" * 60)

    for rel_path, title in FILES:
        full_path = os.path.join(PROJECT_DIR, rel_path)
        if not os.path.exists(full_path):
            print(f"  [SKIP] 跳过（不存在）: {rel_path}")
            continue

        print(f"\n[file] {rel_path} →")
        md_content = read_md(full_path)

        out_name = title.replace("金快查-", "")
        docx_path = os.path.join(OUTPUT_DIR, f"{title}.docx")
        pdf_path  = os.path.join(OUTPUT_DIR, f"{title}.pdf")

        try:
            export_docx(md_content, docx_path, title)
        except Exception as e:
            print(f"  [FAIL] .docx 失败: {e}")

        try:
            export_pdf(md_content, pdf_path, title)
        except Exception as e:
            print(f"  [FAIL] .pdf  失败: {e}")

    print(f"\n{'='*60}")
    print(f"  导出完成 → {OUTPUT_DIR}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
